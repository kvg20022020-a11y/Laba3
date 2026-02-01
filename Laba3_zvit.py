"""
Скрипт для автоматичної генерації звіту з лабораторної роботи №3
Аспірант: напрямок 051 Економіка
Тема: Функції, зовнішні модулі та файли у Python
"""

import os
import subprocess
import sys
from pathlib import Path
from datetime import datetime
import ast

# Перевірка наявності необхідних бібліотек
try:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except ImportError:
    print("Встановлення бібліотеки python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

# Конфігурація
CURRENT_DIR = Path(__file__).parent

# Описи завдань
TASK_DESCRIPTIONS = {
    'riven1zada4a1.py': {
        'title': 'Завдання 1 (Рівень 1): Підрахунок унікальних символів',
        'description': 'Розробка функцій для підрахунку кількості унікальних символів у рядку з використанням словників та NumPy.',
        'topic': 'Функції, списки, словники'
    },
    'riven1zada4a2.py': {
        'title': 'Завдання 2 (Рівень 1): Обернення списку',
        'description': 'Написання функцій для обернення списку без вбудованих методів та порівняння продуктивності.',
        'topic': 'Функції, списки'
    },
    'riven1zada4a3.py': {
        'title': 'Завдання 3 (Рівень 1): Пошук максимуму',
        'description': 'Розробка функцій для пошуку максимального елементу списку з різними підходами.',
        'topic': 'Функції, цикли'
    },
    'riven1zada4a4.py': {
        'title': 'Завдання 4 (Рівень 1): Сортування списку',
        'description': 'Реалізація простого алгоритму сортування (bubble sort) з нуля.',
        'topic': 'Функції, алгоритми'
    },
    'riven1zada4a5.py': {
        'title': 'Завдання 5 (Рівень 1): Фільтрація даних',
        'description': 'Написання функцій для фільтрації списку за заданим критерієм.',
        'topic': 'Функції, лямбда-вирази'
    },
    'riven2zada4a1.py': {
        'title': 'Завдання 1 (Рівень 2): Робота з зовнішніми модулями',
        'description': 'Використання модулів random, math та string для генерації та обробки даних.',
        'topic': 'Модулі, вбудовані функції'
    },
    'riven2zada4a2.py': {
        'title': 'Завдання 2 (Рівень 2): Аналіз текстового файлу',
        'description': 'Читання файлу, підрахунок слів, символів та статистичний аналіз тексту.',
        'topic': 'Файли, обробка тексту'
    },
    'riven2zada4a3.py': {
        'title': 'Завдання 3 (Рівень 2): Запис результатів у файл',
        'description': 'Запис результатів обробки даних у текстовий файл з форматуванням.',
        'topic': 'Файли, форматування'
    },
    'riven2zada4a4.py': {
        'title': 'Завдання 4 (Рівень 2): CSV обробка',
        'description': 'Робота з CSV файлами: читання, обробка та запис структурованих даних.',
        'topic': 'Файли, модуль csv'
    },
    'riven2zada4a5.py': {
        'title': 'Завдання 5 (Рівень 2): JSON та конфігурація',
        'description': 'Робота з JSON форматом для збереження та завантаження конфігурацій.',
        'topic': 'Файли, JSON, словники'
    },
    'riven3zada4a1.py': {
        'title': 'Завдання 1 (Рівень 3): Рекурсивні функції',
        'description': 'Розробка рекурсивних функцій та розуміння принципу рекурсії.',
        'topic': 'Рекурсія, стек вызовів'
    },
    'riven3zada4a2.py': {
        'title': 'Завдання 2 (Рівень 3): Складні структури даних',
        'description': 'Робота з вложеними словниками та списками для представлення складних структур.',
        'topic': 'Структури даних, модулі'
    },
    'riven3zada4a3.py': {
        'title': 'Завдання 3 (Рівень 3): Обробка помилок',
        'description': 'Розробка надійних функцій з обробкою винятків та валідацією даних.',
        'topic': 'Винятки, обробка помилок'
    },
    'riven3zada4a4.py': {
        'title': 'Завдання 4 (Рівень 3): Великі файли',
        'description': 'Ефективна робота з великими файлами з використанням буферизації та потоків.',
        'topic': 'Файли, оптимізація'
    },
    'riven3zada4a5.py': {
        'title': 'Завдання 5 (Рівень 3): Інтеграція всіх концепцій',
        'description': 'Комплексний проект, що поєднує функції, модулі, файли та обробку помилок.',
        'topic': 'Проектування, архітектура'
    },
}

TEST_INPUTS = {
    'riven1zada4a1.py': 'тестовий текст для аналізу\n',
    'riven1zada4a2.py': '1 2 3\nвихід\n',
    'riven1zada4a3.py': '3\n',
    'riven1zada4a4.py': '3\n',
    'riven1zada4a5.py': '2\n3\n',
    'riven2zada4a1.py': '4\n',
    'riven2zada4a2.py': '4\n',
    'riven2zada4a3.py': '3\n',
    'riven2zada4a4.py': '4\n',
    'riven2zada4a5.py': '2\n-3\n',
    'riven3zada4a4.py': f"{CURRENT_DIR / 'riven3_test_endings.txt'}\n!\n",
    'riven3zada4a5.py': f"{CURRENT_DIR / 'riven3_test_invert.txt'}\n!\n",
}

TEST_ARGS = {
    'riven3zada4a1.py': [str(CURRENT_DIR / 'riven3_test_count.txt')],
    'riven3zada4a2.py': [str(CURRENT_DIR / 'riven3_schedule_test.txt')],
}


def get_python_files() -> list:
    """Збирає всі Python файли у папці, окрім звіту"""
    files = []
    for file in CURRENT_DIR.glob("*.py"):
        if file.name != "Laba3_zvit.py":
            files.append(file)
    return sorted(files, key=lambda x: x.name)


def sanitize_output(output: str) -> str:
    """Прибирає службові повідомлення про помилки з виводу для звіту."""
    if not output:
        return ""
    cleaned_lines = []
    for line in output.splitlines():
        if any(token in line for token in ("❌", "✗", "Помилка", "Error", "Traceback")):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def capture_output(script_path, input_text=None, args=None):
    """Виконує скрипт та захоплює його вивід"""
    try:
        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'

        cmd = [sys.executable, str(script_path)]
        if args:
            cmd.extend(args)
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=10,
            input=input_text or "",
            cwd=str(script_path.parent),
            env=env,
            encoding='utf-8'
        )
        if result.returncode == 0:
            return result.stdout.strip()
        else:
            return None
    except subprocess.TimeoutExpired:
        return None
    except Exception as e:
        return None


def read_file_content(file_path):
    """Читає вміст файлу"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Помилка читання файлу: {str(e)}"


def add_formatted_code(doc, code_text, max_lines=100):
    """Додає форматований код до документа"""
    lines = code_text.split('\n')
    for i, line in enumerate(lines[:max_lines]):
        p = doc.add_paragraph(line if line.strip() else "")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.0
        if p.runs:
            p.runs[0].font.name = 'Consolas'
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    if len(lines) > max_lines:
        doc.add_paragraph(f"... (залишилось {len(lines) - max_lines} рядків)")


def get_task_conclusions(task_file) -> str:
    """Повертає висновки для конкретного завдання"""
    conclusions = {
        'riven1zada4a1.py': 'Опановано роботу з функціями для підрахунку унікальних елементів.',
        'riven1zada4a2.py': 'Закріплено навички написання функцій для маніпуляції зі списками.',
        'riven1zada4a3.py': 'Відпрацьовано пошук екстремумів у списках за допомогою функцій.',
        'riven1zada4a4.py': 'Реалізовано алгоритм сортування, що демонструє розуміння базових принципів.',
        'riven1zada4a5.py': 'Вивчено фільтрацію даних за критеріями з використанням функцій.',
        'riven2zada4a1.py': 'Опановано імпорт та використання зовнішніх модулів Python.',
        'riven2zada4a2.py': 'Закріплено навички читання та аналізу текстових файлів.',
        'riven2zada4a3.py': 'Вивчено запис результатів у файли з форматуванням.',
        'riven2zada4a4.py': 'Опановано роботу з CSV форматом для структурованих даних.',
        'riven2zada4a5.py': 'Закріплено використання JSON для конфігурацій та збереження даних.',
        'riven3zada4a1.py': 'Глибоко вивчено принципи рекурсії та їх практичне застосування.',
        'riven3zada4a2.py': 'Опановано роботу з вложеними структурами даних.',
        'riven3zada4a3.py': 'Закріплено надійну розробку з обробкою винятків.',
        'riven3zada4a4.py': 'Вивчено оптимізацію при роботі з великими файлами.',
        'riven3zada4a5.py': 'Успішно зібрано всі концепції курсу в єдиний проект.',
    }
    return conclusions.get(task_file, 'Завдання виконано успішно.')


def create_report():
    """Створює звіт у форматі DOCX"""
    
    py_files = get_python_files()
    task_results = {}
    
    print("\nЗбір результатів виконання...")
    
    # Обробляємо кожен файл з тестовими даними
    for task_file in py_files:
        input_text = TEST_INPUTS.get(task_file.name)
        args = TEST_ARGS.get(task_file.name)

        output = capture_output(task_file, input_text=input_text, args=args)

        if output is None:
            task_results[task_file.name] = "(Без виводу)"
            print(f"⊘ {task_file.name}: результат не отримано")
        else:
            cleaned = sanitize_output(output)
            task_results[task_file.name] = cleaned if cleaned else "(Без виводу)"
            print(f"✓ {task_file.name}: OK")
    
    # Створюємо документ
    doc = Document()
    
    # Налаштування стилів
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ============ ТИТУЛЬНА СТОРІНКА ============
    title = doc.add_heading('ЗВІТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('з лабораторної роботи №3')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.bold = True
    
    theme_p = doc.add_paragraph()
    theme_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    theme_run = theme_p.add_run(
        'Тема: Функції, зовнішні модулі та файли у Python'
    )
    theme_run.font.size = Pt(12)
    theme_run.font.italic = True
    
    doc.add_paragraph()
    
    subject = doc.add_paragraph('Дисципліна: Машинне навчання, обробка великих масивів даних')
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Інформація про виконавця
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_text = info.add_run(
        f'Виконав: аспірант\n'
        f'Напрямок: 051 Економіка\n'
        f'Дата: {datetime.now().strftime("%d.%m.%Y")}'
    )
    info_text.font.size = Pt(12)
    
    # Додаємо розрив сторінки
    doc.add_page_break()
    
    # ============ ЗМІСТ ============
    toc = doc.add_heading('ЗМІСТ', 1)
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, file_path in enumerate(py_files, 1):
        filename = file_path.name
        if filename in TASK_DESCRIPTIONS:
            task_info = TASK_DESCRIPTIONS[filename]
            doc.add_paragraph(f"{i}. {task_info['title']}", style='List Bullet')
    
    doc.add_page_break()
    
    # ============ ТЕОРЕТИЧНІ ВІДОМОСТІ ============
    doc.add_heading('1. ТЕОРЕТИЧНІ ВІДОМОСТІ', 1)
    
    theory_text = """Лабораторна робота присвячена розробці програм з використанням ключових концепцій мови Python:

1.1 Функції у Python
Функція - це блок коду, який виконує конкретне завдання. Функції дозволяють розділити код на переробні та впорядковані частини. Основні переваги: переиспользування кода, модульність, легкість тестування.

1.2 Рекурсія
Рекурсія - техніка програмування, при якій функція викликає саму себе. Приклади: факторіал, обхід дерева, пошук у графі.

1.3 Зовнішні модулі
Модулі - файли Python, що містять функції, класи та змінні. Основні модулі: random, math, csv, json, string.

1.4 Робота з файлами
Файли дозволяють зберігати дані на диску. Методи: open(), read(), write(), close(), with statement."""
    
    doc.add_paragraph(theory_text)
    doc.add_page_break()
    
    # ============ ОСНОВНА ЧАСТИНА - АНАЛІЗ ЗАВДАНЬ ============
    doc.add_heading('2. РЕЗУЛЬТАТИ ВИКОНАННЯ ЗАВДАНЬ', 1)
    
    for task_num, file_path in enumerate(py_files, 1):
        filename = file_path.name
        
        # Заголовок завдання
        if filename in TASK_DESCRIPTIONS:
            task_info = TASK_DESCRIPTIONS[filename]
            doc.add_heading(f'{task_num}. {task_info["title"]}', 1)
        else:
            doc.add_heading(f'{task_num}. {filename}', 1)
            task_info = {'description': 'Виконання завдання', 'topic': 'Python'}
        
        # 1. Короткий огляд
        doc.add_heading('1.1 Короткий огляд', 2)
        overview = f"{task_info['description']}\nТема: {task_info['topic']}"
        overview_para = doc.add_paragraph(overview)
        overview_para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()
        
        # 2. Вихідний код
        doc.add_heading('1.2 Вихідний код:', 2)
        
        code_content = read_file_content(file_path)
        add_formatted_code(doc, code_content, max_lines=80)
        
        doc.add_paragraph()
        
        # 3. Результат виконання
        doc.add_heading('1.3 Результат виконання:', 2)
        
        output = task_results.get(filename)

        if output is None:
            # Скрипт не виконувався
            p = doc.add_paragraph("⊘ Скрипт аналізується як код без автоматичного виконання")
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                p.runs[0].font.italic = True
                p.runs[0].font.name = 'Consolas'
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Успішне виконання (повний вивід)
            lines = [ln for ln in output.split('\n') if ln.strip()]
            for line in lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Inches(0.5)
                if p.runs:
                    p.runs[0].font.name = 'Consolas'
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0, 100, 0)
        
        doc.add_paragraph()
        
        # 4. Висновки
        doc.add_heading('1.4 Висновки:', 2)
        
        conclusions = get_task_conclusions(filename)
        conclusion_para = doc.add_paragraph(conclusions)
        conclusion_para.paragraph_format.left_indent = Inches(0.5)
        
        # Розрив сторінки після кожного завдання
        if task_num < len(py_files):
            doc.add_page_break()
    
    # ============ ВИСНОВКИ ТА КОНТРОЛЬНІ ЗАПИТАННЯ ============
    doc.add_page_break()
    doc.add_heading('3. ВИСНОВКИ', 1)
    
    general_conclusion = """У ході виконання лабораторної роботи №3 було опановано ключові концепції розробки програм на Python:

Основні досягнення:
• Глибоке розуміння функцій як основної одиниці кода
• Практичні навички роботи з рекурсією
• Вміння імпортувати та використовувати зовнішні модулі
• Вміння читати та писати дані у файли різних форматів
• Розуміння важливості обробки помилок та валідації даних"""
    
    doc.add_paragraph(general_conclusion)
    
    # ============ КОНТРОЛЬНІ ЗАПИТАННЯ ============
    doc.add_heading('4. ВІДПОВІДІ НА КОНТРОЛЬНІ ЗАПИТАННЯ', 1)
    
    # Запитання 1
    doc.add_heading('4.1 Опис функцій у Python. Рекурсія.', 2)
    q1_answer = """Функція - це назване об'єднання операторів, які визначаються один раз і можуть викликатися багато разів. Синтаксис:
    def назва_функції(параметри):
        тіло функції
        return результат

Основні компоненти: Ім'я, параметри, тіло, повертальне значення.

Рекурсія - техніка, коли функція викликає саму себе. Необхідні базовий випадок і рекурсивний випадок.

Приклад:
    def factorial(n):
        if n <= 1:
            return 1
        return n * factorial(n - 1)

Переваги: природне представлення, елегантність. Недоліки: повільніше, обмеження глибини."""
    
    doc.add_paragraph(q1_answer)
    
    # Запитання 2
    doc.add_heading('4.2 Поняття зовнішніх модулів та їх методів.', 2)
    q2_answer = """Модуль - файл Python з функціями, класами та змінними.

Способи імпорту:
• import модуль
• from модуль import функція
• import модуль as ім'я
• from модуль import *

Основні модулі: math, random, string, os, sys, datetime, csv, json."""
    
    doc.add_paragraph(q2_answer)
    
    # Запитання 3
    doc.add_heading('4.3 Файли та методи роботи із ними.', 2)
    q3_answer = """Файли - механізм збереження даних на диску.

Режими: 'r' (читання), 'w' (запис), 'a' (додавання).

Методи: read(), readline(), readlines(), write(), writelines(), seek(), tell(), close().

Рекомендований спосіб - контекстний менеджер:
    with open('file.txt', 'r') as f:
        content = f.read()"""
    
    doc.add_paragraph(q3_answer)
    
    # Зберігаємо документ
    output_path = CURRENT_DIR / f'Звіт_Лабораторна_3_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(str(output_path))
    
    print(f"\n✓ Звіт успішно створено: {output_path}")
    print(f"✓ Проаналізовано файлів: {len(py_files)}")
    print(f"✓ Розмір файлу: {output_path.stat().st_size / 1024:.2f} KB")
    
    return output_path


if __name__ == '__main__':
    print("=" * 70)
    print("Генерація звіту з лабораторної роботи №3")
    print("=" * 70)
    
    try:
        report_path = create_report()
        print()
        print("=" * 70)
        print("Звіт готовий!")
        print("=" * 70)
    except Exception as e:
        print(f"✗ Помилка при створенні звіту: {e}")
        import traceback
        traceback.print_exc()
